import os
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from docx.shared import RGBColor as DocxRGBColor, Inches as DocxInches
from pptx import Presentation
from pptx.util import Inches, Pt
from pptx.enum.shapes import MSO_SHAPE
from pptx.dml.color import RGBColor as PptxRGBColor
from tempfile import NamedTemporaryFile
from PIL import Image, ImageDraw, ImageFont
from docx2pdf import convert as docx2pdf_convert
import shutil

DPI = 96  # Standard screen DPI for conversion

def generate_watermark_image(text, width=600, height=200, color=(80, 80, 80), opacity=80, font_size=48, position='Center Diagonal'):
    image = Image.new('RGBA', (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype('arial.ttf', font_size)
    except:
        font = ImageFont.load_default()
    if hasattr(draw, 'textbbox'):
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
    else:
        text_width, text_height = font.getsize(text)
    # Default position: center
    x = (width - text_width) / 2
    y = (height - text_height) / 2
    if position == 'Top-left':
        x, y = 10, 10
    elif position == 'Top-right':
        x, y = width - text_width - 10, 10
    elif position == 'Bottom-left':
        x, y = 10, height - text_height - 10
    elif position == 'Bottom-right':
        x, y = width - text_width - 10, height - text_height - 10
    elif position == 'Center':
        x, y = (width - text_width) / 2, (height - text_height) / 2
    # Draw text (with or without rotation)
    txt_img = Image.new('RGBA', (width, height), (255, 255, 255, 0))
    txt_draw = ImageDraw.Draw(txt_img)
    txt_draw.text((x, y), text, font=font, fill=(color[0], color[1], color[2], int(opacity/100*255)))
    if position == 'Center Diagonal':
        rotated = txt_img.rotate(30, expand=1)
        image.alpha_composite(rotated, (0, 0))
    else:
        image.alpha_composite(txt_img, (0, 0))
    temp = NamedTemporaryFile(delete=False, suffix='.png')
    image.save(temp.name, 'PNG')
    return temp.name

def add_watermark_pdf(input_path, watermark_text, output_path, color=(80,80,80), opacity=80, position='Center Diagonal', font_size=48, progress_callback=None):
    reader = PdfReader(input_path)
    writer = PdfWriter()
    from io import BytesIO
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.colors import Color
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    can.setFont("Helvetica", font_size)
    r, g, b = [c/255 for c in color]
    can.setFillColor(Color(r, g, b, alpha=opacity/100))
    can.saveState()
    tx, ty = 300, 400
    if position == 'Top-left':
        tx, ty = 100, 700
    elif position == 'Top-right':
        tx, ty = 500, 700
    elif position == 'Bottom-left':
        tx, ty = 100, 100
    elif position == 'Bottom-right':
        tx, ty = 500, 100
    elif position == 'Center':
        tx, ty = 300, 400
    can.translate(tx, ty)
    if position == 'Center Diagonal':
        can.rotate(45)
    can.drawCentredString(0, 0, watermark_text)
    can.restoreState()
    can.save()
    packet.seek(0)
    from PyPDF2 import PdfReader as RLReader
    watermark_pdf = RLReader(packet)
    watermark_page = watermark_pdf.pages[0]
    total_pages = len(reader.pages)
    for i, page in enumerate(reader.pages):
        page.merge_page(watermark_page)
        writer.add_page(page)
        if progress_callback:
            progress_callback(i + 1, total_pages)
    with open(output_path, 'wb') as f:
        writer.write(f)

def add_watermark_docx(input_path, watermark_text, output_path, color=(80,80,80), opacity=80, position='Center Diagonal', font_size=48):
    doc = Document(input_path)
    for section in doc.sections:
        page_width_in = section.page_width / 914400  # EMU to inches
        page_height_in = section.page_height / 914400
        img_width_px = int(page_width_in * DPI)
        img_height_px = int(page_height_in * 0.15 * DPI)  # 15% of page height
        watermark_img = generate_watermark_image(watermark_text, width=img_width_px, height=img_height_px, color=color, opacity=opacity, font_size=font_size, position=position)
        section.header_distance = DocxInches(1.5)
        header = section.header
        for shape in header._element.xpath('.//w:drawing'):
            shape.getparent().remove(shape)
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(watermark_img, width=DocxInches(page_width_in))
        paragraph.alignment = 1  # Center
        os.unlink(watermark_img)
    doc.save(output_path)

def encrypt_pdf(input_pdf, output_pdf, password):
    reader = PdfReader(input_pdf)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    with open(output_pdf, 'wb') as f:
        writer.write(f)

def add_watermark(input_path, watermark_text, output_path, color=(80,80,80), opacity=80, position='Center Diagonal', font_size=48, password=None, progress_callback=None):
    ext = os.path.splitext(input_path)[1].lower()
    if ext == '.pdf':
        temp_out = output_path if not password else output_path + '.tmp.pdf'
        add_watermark_pdf(input_path, watermark_text, temp_out, color=color, opacity=opacity, position=position, font_size=font_size, progress_callback=progress_callback)
        if password:
            encrypt_pdf(temp_out, output_path, password)
            if os.path.exists(temp_out):
                os.remove(temp_out)
    elif ext == '.docx':
        with NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            temp_pdf_path = temp_pdf.name
        try:
            docx2pdf_convert(input_path, temp_pdf_path)
            temp_out = output_path if not password else output_path + '.tmp.pdf'
            add_watermark_pdf(temp_pdf_path, watermark_text, temp_out, color=color, opacity=opacity, position=position, font_size=font_size, progress_callback=progress_callback)
            if password:
                encrypt_pdf(temp_out, output_path, password)
                if os.path.exists(temp_out):
                    os.remove(temp_out)
        finally:
            if os.path.exists(temp_pdf_path):
                os.remove(temp_pdf_path)
    else:
        raise ValueError('Only PDF and Word (.docx) files are supported.') 